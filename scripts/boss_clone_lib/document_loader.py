"""アップロードされた資料からテキストを抽出する。

仕様: Day 4 後半 ファイル添付機能 Phase 1
対応形式: .txt / .md / .pdf / .docx
50,000 字を超えたら警告付きで切り詰める。

Phase 2（出張明け検討）で .png / .jpg 等の Vision 対応を入れる想定。
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Any

MAX_CHARS = 50_000

SUPPORTED_FORMATS = {"txt", "md", "pdf", "docx"}


def _detect_format(filename: str) -> str:
    suffix = Path(filename or "").suffix.lower().lstrip(".")
    if suffix in {"txt"}:
        return "txt"
    if suffix in {"md", "markdown"}:
        return "md"
    if suffix == "pdf":
        return "pdf"
    if suffix == "docx":
        return "docx"
    return suffix or "unknown"


def _extract_pdf(data: bytes) -> str:
    """pypdf で全ページからテキスト抽出（失敗ページはスキップして続行）。"""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            t = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            t = f"[page {i+1}: テキスト抽出失敗]"
        if t.strip():
            parts.append(t.strip())
    return "\n\n".join(parts)


def _extract_docx(data: bytes) -> str:
    """python-docx で paragraphs + table cells テキスト抽出。"""
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
    # テーブル本文も拾う（PdM 資料はテーブル多用）
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join((cell.text or "").strip() for cell in row.cells if (cell.text or "").strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _extract_text(data: bytes) -> str:
    """テキスト系ファイル。UTF-8 → cp932 → latin-1 のフォールバック。"""
    for enc in ("utf-8", "utf-8-sig", "cp932", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def load_document(filename: str, data: bytes) -> dict[str, Any]:
    """ファイル名 + バイト列から内容を抽出。

    返り値:
        {
            "filename": str,
            "format": "pdf|docx|txt|md|unknown",
            "content": str,        # 抽出テキスト（必要に応じて切り詰め済み）
            "char_count": int,     # 切り詰め後の文字数
            "original_char_count": int,  # 切り詰め前の文字数
            "truncated": bool,     # 50K 超で切り詰めたか
            "warnings": list[str], # ユーザー向け警告メッセージ
            "status": "ok" | "error",
            "error": str | None,
        }
    """
    fmt = _detect_format(filename)
    result: dict[str, Any] = {
        "filename": filename,
        "format": fmt,
        "content": "",
        "char_count": 0,
        "original_char_count": 0,
        "truncated": False,
        "warnings": [],
        "status": "ok",
        "error": None,
    }

    if fmt not in SUPPORTED_FORMATS:
        result["status"] = "error"
        result["error"] = (
            f"非対応の形式: .{fmt}。対応形式: {', '.join('.' + f for f in sorted(SUPPORTED_FORMATS))}"
        )
        return result

    try:
        if fmt == "pdf":
            content = _extract_pdf(data)
        elif fmt == "docx":
            content = _extract_docx(data)
        else:  # txt / md
            content = _extract_text(data)
    except Exception as e:  # noqa: BLE001
        result["status"] = "error"
        result["error"] = f"{type(e).__name__}: {e}"
        return result

    original_len = len(content)
    if original_len > MAX_CHARS:
        content = content[:MAX_CHARS]
        result["truncated"] = True
        result["warnings"].append(
            f"資料が {original_len:,} 字あったので、先頭 {MAX_CHARS:,} 字に切り詰めました。"
            f"後半の内容はレビュー対象外になります。"
        )
    if not content.strip():
        result["warnings"].append("資料からテキストが抽出できませんでした（画像のみの PDF などの可能性）。")

    result["content"] = content
    result["char_count"] = len(content)
    result["original_char_count"] = original_len
    return result


def format_for_prompt(doc: dict, max_chars_in_prompt: int = 8000) -> str:
    """エージェントのプロンプトに埋め込む形式へ整形。

    プロンプト全体のトークン数を抑えるため、デフォルトでは 8K 字に切り詰める。
    """
    content = (doc.get("content") or "")[:max_chars_in_prompt]
    fname = doc.get("filename", "(no name)")
    n = doc.get("char_count", 0)
    fmt = doc.get("format", "?")
    excerpt_note = ""
    if doc.get("original_char_count", 0) > max_chars_in_prompt:
        excerpt_note = f"（※ 全 {doc['original_char_count']:,} 字のうち先頭 {len(content):,} 字を抜粋）"
    return (
        f"ファイル名: {fname}\n"
        f"形式: .{fmt} / 抽出 {n:,} 字 {excerpt_note}\n"
        f"---\n"
        f"{content}\n"
        f"---"
    )
